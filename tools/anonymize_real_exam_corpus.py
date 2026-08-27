"""Create a publishable MAHİR regression corpus from the private test ZIP."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import unicodedata
from collections import defaultdict
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile

from docx import Document
from PIL import Image, ImageDraw, ImageFont


ANON_TEXT = {
    "erzurum": "TEST İLİ",
    "yakutiye": "TEST İLÇESİ",
    "mahir anadolu lisesi": "MAHİR TEST LİSESİ",
    "mahır anadolu lisesi": "MAHİR TEST LİSESİ",
    "zülal ülker daştan": "TEST ÖĞRETMENİ",
}


def normalise(value: object) -> str:
    text = str(value or "").strip().lower().replace("ı", "i")
    return "".join(char for char in unicodedata.normalize("NFKD", text) if not unicodedata.combining(char))


def replace_paragraph_text(paragraph) -> None:
    if not paragraph.text:
        return
    updated = paragraph.text
    for original, replacement in ANON_TEXT.items():
        updated = re.sub(re.escape(original), replacement, updated, flags=re.IGNORECASE)
    if updated != paragraph.text:
        paragraph.text = updated


def anonymise_docx(source: Path, destination: Path) -> None:
    document = Document(source)
    for paragraph in document.paragraphs:
        replace_paragraph_text(paragraph)
    for section in document.sections:
        for story in (section.header, section.footer):
            for paragraph in story.paragraphs:
                replace_paragraph_text(paragraph)

    student_index = 0
    for table in document.tables:
        for row in table.rows:
            cells = row.cells
            for index, cell in enumerate(cells[:-1]):
                label = normalise(cell.text)
                if label in {"il", "il:"}:
                    cells[index + 1].text = "TEST İLİ"
                elif label in {"ilce", "ilce:"}:
                    cells[index + 1].text = "TEST İLÇESİ"
                elif "okul adi" in label or "okul/kurum adi" in label:
                    cells[index + 1].text = "MAHİR TEST LİSESİ"
                elif "ogretmenin adi soyadi" in label or "ogretmen adi soyadi" in label:
                    cells[index + 1].text = "TEST ÖĞRETMENİ"
                elif "ogrencinin adi-soyadi" in label or "ogrencinin adi soyadi" in label:
                    cells[index + 1].text = "ÖĞRENCİ-001"
                elif "ogrenci okul no" in label:
                    cells[index + 1].text = "OGR-001"
            for cell in cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_text(paragraph)

        if not table.rows:
            continue
        headings = [normalise(cell.text) for cell in table.rows[0].cells]
        number_index = next(
            (index for index, heading in enumerate(headings) if heading in {"okul no", "ogrenci no", "ogrenci numarasi"}),
            None,
        )
        if number_index is None:
            continue
        for row in table.rows[1:]:
            if number_index >= len(row.cells):
                continue
            current = row.cells[number_index].text.strip()
            if not current or normalise(current) in {"azami", "azami puan", "-"}:
                continue
            student_index += 1
            row.cells[number_index].text = f"OGR-{student_index:03d}"

    document.core_properties.author = "MAHİR TEST"
    document.core_properties.last_modified_by = "MAHİR TEST"
    document.core_properties.comments = "Anonim MAHİR regresyon evrakı"
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def load_font(size: int):
    for candidate in (Path("C:/Windows/Fonts/segoeui.ttf"), Path("C:/Windows/Fonts/arial.ttf")):
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def save_clean_image(source: Path, destination: Path, *, comparison_index: int | None = None) -> None:
    with Image.open(source) as opened:
        image = opened.convert("RGB")
    if image.width > 2400:
        ratio = 2400 / image.width
        image = image.resize((2400, round(image.height * ratio)), Image.Resampling.LANCZOS)
    if comparison_index is not None:
        draw = ImageDraw.Draw(image)
        width, height = image.size
        # El yazısı hücrenin sol kenarından başladığı için sınırı dikey tablo
        # çizgisine kadar kapat; ilk birkaç harfin dışarıda kalmasına izin verme.
        x1, x2 = round(width * 0.33), round(width * 0.985)
        name_y1, name_y2 = round(height * 0.012), round(height * 0.122)
        no_y1, no_y2 = round(height * 0.142), round(height * 0.252)
        draw.rectangle((x1, name_y1, x2, name_y2), fill="white")
        draw.rectangle((x1, no_y1, x2, no_y2), fill="white")
        text_font = load_font(max(18, round(height * 0.075)))
        draw.text((x1 + 14, name_y1 + 3), f"ÖĞRENCİ-{comparison_index:03d}", fill="#123a8c", font=text_font)
        draw.text((x1 + 14, no_y1 + 3), f"OGR-{comparison_index:03d}", fill="#123a8c", font=text_font)
    destination.parent.mkdir(parents=True, exist_ok=True)
    if destination.suffix.lower() == ".png":
        image.save(destination, format="PNG", optimize=True)
    else:
        image.save(destination, format="JPEG", quality=88, optimize=True, progressive=True)


def classify_png_group(files: list[Path]) -> str:
    stems = [file.stem.lower() for file in files]
    if len(files) == 25:
        return "images/9A-written-25"
    if len(files) == 20 and all(stem.endswith("d") for stem in stems):
        return "images/9B-listening-20"
    if len(files) == 20 and all(stem.endswith("k") for stem in stems):
        return "images/9B-speaking-20"
    if len(files) == 20 and all(stem.endswith("y") for stem in stems):
        return "images/9B-written-20"
    raise ValueError(f"Tanınmayan PNG grubu: {len(files)} dosya")


def classify_docx_groups(files: list[Path]) -> dict[Path, str]:
    by_parent: dict[Path, list[Path]] = defaultdict(list)
    for file in files:
        by_parent[file.parent].append(file)
    destinations: dict[Path, str] = {}
    for parent_files in by_parent.values():
        names = [file.name.lower() for file in parent_files]
        if len(parent_files) == 5 and all("yazili_puan_cizelgesi" in name for name in names):
            folder = "word/five-classes-written"
        elif len(parent_files) == 5 and all("dinleme_puan_cizelgesi" in name for name in names):
            folder = "word/five-classes-listening"
        elif len(parent_files) == 5 and all("konusma_puan_cizelgesi" in name for name in names):
            folder = "word/five-classes-speaking"
        elif any("ogrenme_ciktilari" in name for name in names):
            folder = "word/shared-learning-outcomes"
        elif len(parent_files) == 3:
            folder = "word/ocr-word-comparison"
        else:
            folder = "word/templates"
        for file in parent_files:
            destinations[file] = folder
    return destinations


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def build(source_zip: Path, output: Path) -> None:
    if output.exists() and any(output.iterdir()):
        raise FileExistsError(f"Çıktı klasörü boş olmalıdır: {output}")
    output.mkdir(parents=True, exist_ok=True)
    with TemporaryDirectory(prefix="mahir-anon-source-") as temporary:
        source_root = Path(temporary)
        with ZipFile(source_zip) as archive:
            archive.extractall(source_root)
        source_files = [path for path in source_root.rglob("*") if path.is_file()]
        png_groups: dict[Path, list[Path]] = defaultdict(list)
        for path in source_files:
            if path.suffix.lower() == ".png":
                png_groups[path.parent].append(path)
        for files in png_groups.values():
            folder = output / classify_png_group(files)
            for index, source in enumerate(sorted(files), start=1):
                save_clean_image(source, folder / f"exam-{index:03d}.png")
        jpeg_files = sorted(path for path in source_files if path.suffix.lower() in {".jpg", ".jpeg"})
        for index, source in enumerate(jpeg_files, start=1):
            save_clean_image(source, output / "images/9A-ocr-word-comparison-25" / f"exam-{index:03d}.jpeg", comparison_index=index)
        docx_files = [path for path in source_files if path.suffix.lower() == ".docx"]
        destinations = classify_docx_groups(docx_files)
        used_names: dict[str, int] = defaultdict(int)
        for source in sorted(docx_files):
            folder = destinations[source]
            stem = re.sub(r"[^A-Za-z0-9_-]+", "_", source.stem).strip("_") or "document"
            key = f"{folder}/{stem}"
            used_names[key] += 1
            suffix = f"-{used_names[key]}" if used_names[key] > 1 else ""
            anonymise_docx(source, output / folder / f"{stem}{suffix}.docx")
    files = sorted(path for path in output.rglob("*") if path.is_file())
    manifest = {
        "privacy": "All identities and document metadata are synthetic.",
        "counts": {
            "png": sum(path.suffix.lower() == ".png" for path in files),
            "jpeg": sum(path.suffix.lower() in {".jpg", ".jpeg"} for path in files),
            "docx": sum(path.suffix.lower() == ".docx" for path in files),
        },
        "generalEvaluationWeights": {"written": 70, "listening": 15, "speaking": 15},
        "files": [
            {"path": path.relative_to(output).as_posix(), "sha256": sha256(path), "bytes": path.stat().st_size}
            for path in files
        ],
    }
    (output / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_zip", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.source_zip, args.output)


if __name__ == "__main__":
    main()
