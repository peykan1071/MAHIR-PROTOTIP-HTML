"""Regression tests for the two official MAHIR Word score templates."""

from __future__ import annotations

import io
import unittest
from pathlib import Path

from docx import Document

from backend.app.docx_parser import parse_mahir_docx


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = ROOT / "shared" / "templates"


class OfficialDocxTemplateTests(unittest.TestCase):
    def parse_template(self, name: str) -> dict[str, object]:
        return parse_mahir_docx((TEMPLATES / name).read_bytes())

    def filled_template(self, name: str, fill) -> dict[str, object]:
        document = Document(TEMPLATES / name)
        fill(document)
        output = io.BytesIO()
        document.save(output)
        return parse_mahir_docx(output.getvalue())

    def test_class_template_is_read_without_ocr_or_invented_outcomes(self):
        result = self.parse_template("MAHIR_Veri_Giris_Sablonu.docx")

        self.assertEqual(result["documentType"], "mahir-class-score-template")
        self.assertEqual(result["summary"]["questionCount"], 10)
        self.assertEqual(result["summary"]["studentCount"], 0)
        self.assertEqual(
            [question["number"] for question in result["questions"]],
            list(range(1, 11)),
        )
        self.assertTrue(all(question["maxScore"] is None for question in result["questions"]))
        self.assertTrue(all(not question["outcomeCode"] for question in result["questions"]))
        self.assertIn(
            "Öğrenme çıktıları puan çizelgesinde yer almıyor",
            " ".join(result["warnings"]),
        )
        self.assertNotIn("ad-soyad", " ".join(result["students"]))

    def test_single_student_sheet_reads_question_maxima_without_identity_data(self):
        result = self.parse_template("Ornek_Sinav_Kagidi_Soru_Bazli_Puan_Cizelgesi.docx")

        self.assertEqual(result["documentType"], "single-student-score-sheet")
        self.assertEqual(result["summary"]["questionCount"], 10)
        self.assertEqual(result["summary"]["studentCount"], 0)
        self.assertEqual(
            [question["maxScore"] for question in result["questions"]],
            [10] * 10,
        )
        self.assertIn("ad-soyad", " ".join(result["warnings"]))
        self.assertNotIn(
            "öğrenci ve puan alanları otomatik olarak ayırt edilemedi",
            " ".join(result["warnings"]),
        )

    def test_filled_class_template_reads_school_number_scores_and_total(self):
        def fill(document: Document) -> None:
            row = document.tables[1].rows[1].cells
            row[1].text = "101"
            for index in range(2, 12):
                row[index].text = "10"
            row[12].text = "100"

        result = self.filled_template("MAHIR_Veri_Giris_Sablonu.docx", fill)

        self.assertEqual(result["summary"]["studentCount"], 1)
        self.assertEqual(result["students"][0]["studentNo"], "101")
        self.assertEqual(result["students"][0]["scores"], [10] * 10)
        self.assertEqual(result["students"][0]["calculatedTotal"], 100)

    def test_fixed_ten_question_headings_use_only_active_maximum_columns(self):
        document = Document()
        metadata = document.add_table(rows=2, cols=4)
        for cell, value in zip(
            metadata.rows[0].cells,
            ["Sınıf/Şube", "9-A", "Dersin Adı", "Türk Dili ve Edebiyatı"],
        ):
            cell.text = value
        for cell, value in zip(
            metadata.rows[1].cells,
            ["Sınav Türü", "☒ 2. Dinleme", "Eğitim Öğretim Yılı", "2025-2026"],
        ):
            cell.text = value

        scores = document.add_table(rows=3, cols=13)
        for cell, value in zip(
            scores.rows[0].cells,
            ["Sıra", "Okul No", *[f"Soru {index}" for index in range(1, 11)], "Toplam"],
        ):
            cell.text = value
        for cell, value in zip(
            scores.rows[1].cells,
            ["AZAMİ", "AZAMİ", "25", "25", "25", "25", "-", "-", "-", "-", "-", "-", "100"],
        ):
            cell.text = value
        maximum_label = scores.rows[1].cells[0].merge(scores.rows[1].cells[1])
        maximum_label.text = "AZAMİ"
        for cell, value in zip(
            scores.rows[2].cells,
            ["1", "1001", "22", "20", "15", "16", "-", "-", "-", "-", "-", "-", "73"],
        ):
            cell.text = value

        output = io.BytesIO()
        document.save(output)
        result = parse_mahir_docx(output.getvalue())

        self.assertEqual(result["summary"]["questionCount"], 4)
        self.assertEqual([question["maxScore"] for question in result["questions"]], [25] * 4)
        self.assertEqual(result["students"][0]["scores"], [22, 20, 15, 16])
        self.assertEqual(result["students"][0]["calculatedTotal"], 73)

    def test_filled_single_student_sheet_omits_name_and_reads_scores(self):
        def fill(document: Document) -> None:
            metadata = document.tables[0]
            metadata.rows[0].cells[1].text = "KİŞİSEL AD YAZILSA BİLE ALINMAZ"
            metadata.rows[1].cells[1].text = "202"
            score_row = document.tables[1].rows[2].cells
            for index in range(1, 11):
                score_row[index].text = "8"
            score_row[11].text = "80"

        result = self.filled_template(
            "Ornek_Sinav_Kagidi_Soru_Bazli_Puan_Cizelgesi.docx", fill
        )

        self.assertEqual(result["summary"]["studentCount"], 1)
        self.assertEqual(result["students"][0]["studentNo"], "202")
        self.assertEqual(result["students"][0]["scores"], [8] * 10)
        self.assertEqual(result["students"][0]["totalScore"], 80)
        self.assertNotIn("KİŞİSEL AD", repr(result["students"]))

    def test_compact_simulation_tables_preserve_metadata_outcomes_and_maxima(self):
        document = Document()
        metadata = document.add_table(rows=3, cols=4)
        metadata.rows[0].cells[0].text = "İl / İlçe"
        metadata.rows[0].cells[1].text = "Erzurum / Yakutiye"
        metadata.rows[0].cells[2].text = "Okul"
        metadata.rows[0].cells[3].text = "Örnek Lisesi"
        metadata.rows[1].cells[0].text = "Eğitim yılı"
        metadata.rows[1].cells[1].text = "2025-2026"
        metadata.rows[1].cells[2].text = "Sınıf / Ders"
        metadata.rows[1].cells[3].text = "9-A / Türk Dili ve Edebiyatı"
        metadata.rows[2].cells[0].text = "Sınav türü"
        metadata.rows[2].cells[1].text = "2. Dinleme"

        questions = document.add_table(rows=2, cols=5)
        for cell, value in zip(
            questions.rows[0].cells,
            ["Soru", "Azami", "Kod", "Öğrenme Çıktısı", "Değerlendirme kanıtı"],
        ):
            cell.text = value
        for cell, value in zip(
            questions.rows[1].cells,
            ["S1", "10", "TDE1.1.1", "Seçim yapar.", "Öğretmen gözlemi"],
        ):
            cell.text = value

        students = document.add_table(rows=3, cols=14)
        for cell, value in zip(
            students.rows[0].cells,
            ["Sıra", "Okul No", "S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8", "S9", "S10", "Toplam", "Kontrol"],
        ):
            cell.text = value
        for cell, value in zip(students.rows[1].cells, ["AZAMİ", "—", "10", "", "", "", "", "", "", "", "", "", "10", ""]):
            cell.text = value
        for cell, value in zip(students.rows[2].cells, ["1", "OGR-001", "8", "", "", "", "", "", "", "", "", "", "8", ""]):
            cell.text = value

        output = io.BytesIO()
        document.save(output)
        result = parse_mahir_docx(output.getvalue())

        self.assertEqual(result["exam"]["province"], "Erzurum")
        self.assertEqual(result["exam"]["district"], "Yakutiye")
        self.assertEqual(result["exam"]["classSection"], "9-A")
        self.assertEqual(result["exam"]["course"], "Türk Dili ve Edebiyatı")
        self.assertEqual(result["exam"]["examType"], "Dinleme")
        self.assertEqual(result["questions"][0]["outcomeCode"], "TDE1.1.1")
        self.assertEqual(result["questions"][0]["outcomeDescription"], "Seçim yapar.")
        self.assertEqual(result["questions"][0]["maxScore"], 10)
        self.assertEqual(result["summary"]["studentCount"], 1)
        self.assertEqual(result["students"][0]["studentNo"], "OGR-001")
        self.assertEqual(result["students"][0]["scores"], [8])


if __name__ == "__main__":
    unittest.main()
